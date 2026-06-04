#!/usr/bin/env python3
"""Fill Vibook graduation template Introduction (1.1–1.11) in a Word document."""

import argparse
import re
import shutil
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

DEFAULT_SOURCE = Path(
    "/Users/ayazidan/Downloads/Graduation Project TEMPLATE (1)-2.docx"
)
DEFAULT_OUTPUT = Path("Vibook_Introduction_FIXED.docx")


def _style_ref_cache(doc):
    """Map style names to a paragraph element that already uses that style."""
    cache = {}
    for p in doc.paragraphs:
        name = p.style.name if p.style else None
        if name and name not in cache:
            cache[name] = p._p
    return cache


def _clone_paragraph_element(ref_p):
    """Clone an existing w:p so Word receives valid pPr/numPr/rPr structure."""
    new_p = deepcopy(ref_p)
    for child in list(new_p):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("r", "hyperlink", "smartTag", "ins", "del", "commentRangeStart", "commentRangeEnd", "bookmarkStart", "bookmarkEnd"):
            new_p.remove(child)
    return new_p


def insert_paragraph_after(paragraph, text="", style=None, style_refs=None):
    ref = paragraph._p
    if style and style_refs and style in style_refs:
        ref = style_refs[style]
    new_p = _clone_paragraph_element(ref)
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        paragraph.add_run("")


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def set_paragraph_text(paragraph, text, style=None):
    clear_paragraph(paragraph)
    paragraph.add_run(text)
    if style:
        paragraph.style = style


def find_chapter1_bounds(doc):
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t == "Chapter 1" and start is None:
            start = i
        if t == "Chapter 2" and start is not None:
            end = i
            break
    return start, end


def update_timeline_table(doc):
    """Table 0: project timeline in Chapter 1."""
    if not doc.tables:
        return
    table = doc.tables[0]
    rows = [
        ("WBS", "Task description", "Start date", "Finish date", "Progress"),
        ("1", "Requirements analysis and domain modeling for event booking in Jordan", "Sep 2024", "Oct 2024", "100%"),
        ("2", "System architecture and database design (MySQL, REST API)", "Oct 2024", "Nov 2024", "100%"),
        ("3", "Backend implementation (Spring Boot, JWT, JPA entities)", "Nov 2024", "Feb 2025", "100%"),
        ("4", "Mobile app implementation (Expo, consumer and partner flows)", "Jan 2025", "Apr 2025", "100%"),
        ("5", "Admin web console and integration testing", "Mar 2025", "May 2025", "100%"),
        ("6", "Documentation, UAT, and graduation report preparation", "May 2025", "Jun 2025", "90%"),
    ]
    for ri, row_data in enumerate(rows):
        if ri >= len(table.rows):
            break
        for ci, val in enumerate(row_data):
            if ci < len(table.rows[ri].cells):
                table.rows[ri].cells[ci].text = val


SECTIONS = []  # populated in main


def build_sections():
    return [
        (
            "Project Description",
            "s70",
            (
                "Vibook is a full-stack lifestyle event booking platform developed as a graduation "
                "project in software engineering. The system connects consumers in Jordan with "
                "bookable business events while enabling approved business partners to publish "
                "and manage listings and bookings. A separate administrative web console supports "
                "platform operators in moderating businesses, events, bookings, ratings, and user "
                "reports.\n\n"
                "The project addresses the need for a unified, trustworthy marketplace for timed "
                "experiences: users can discover events by governorate and category, view details, "
                "favorite listings, submit bookings, rate attended events, and report inappropriate "
                "content. Business owners apply for a business profile, await administrative approval, "
                "then create and manage events and incoming reservations. Administrators review "
                "applications, moderate catalog content, resolve reports, and view operational analytics.\n\n"
                "Primary goals include delivering a secure JWT-authenticated REST API backed by MySQL, "
                "a cross-platform mobile client for guests and registered users, and a role-protected "
                "admin SPA. Success is measured by end-to-end implementation of catalog discovery, "
                "booking lifecycle, partner onboarding, moderation workflows, and stable integration "
                "between the three deployable components documented in the project repository."
            ),
        ),
        (
            "Project Overview",
            "s70",
            (
                "The Vibook solution comprises three implemented parts: a Spring Boot backend API "
                "(Java 21, port 8080, prefix /api/v1), an Expo/React Native mobile application "
                "(file-based routing via expo-router), and a Vite/React admin dashboard for users "
                "with the ADMIN role. Persistent data is stored in MySQL database vibook_db using "
                "Hibernate JPA with sixteen entity types including User, BusinessProfile, BusinessEvent, "
                "Booking, Favorite, EventRating, ModerationReport, and reference data for governorates "
                "and categories.\n\n"
                "[PLACEHOLDER: Figure 1.1 - Vibook System Architecture Diagram]\n"
                "Description: Diagram the three clients (mobile, admin-web), Spring Boot API, JWT "
                "security filter, and MySQL vibook_db with major REST route groups (/auth, /events, "
                "/bookings, /business-profile, /business/events, /admin) as implemented in the repository.\n\n"
                "Target users are guests (browse without login), registered consumers (USER role), "
                "business partners (USER with approved BusinessProfile and optional ROLE_BUSINESS), "
                "and platform administrators. Core implemented capabilities include public and authenticated "
                "event catalog access, consumer bookings with PayPal Sandbox payment integration, "
                "favorites, event ratings, user moderation reports, business profile application and approval, "
                "partner event CRUD with photo upload, partner booking status updates, and admin analytics "
                "summary. Presentation-only mobile screens (premium membership, wallet, vouchers, legacy "
                "restaurant/stay routes) are outside the production data scope and are not described as "
                "backend-backed features in this document."
            ),
        ),
        (
            "Tasks",
            "s70",
            (
                "The project was decomposed into verifiable engineering tasks aligned with the repository "
                "structure. Analysis and modeling tasks covered stakeholder needs, governorate/category "
                "taxonomy, and booking state transitions. Backend tasks included security configuration, "
                "JWT access and refresh tokens, twenty-three REST controllers, service-layer business rules, "
                "and JPA repositories with admin list specifications. Mobile tasks covered authentication "
                "hydration, Explore and search flows, event product-detail pages, checkout and payment, "
                "favorites, ratings, reports, Arabic/English localization with RTL layout, and the business "
                "partner hub under app/business/. Admin-web tasks included protected routing, Axios API client, "
                "and pages for businesses, events, bookings, ratings, reports, and dashboard charts.\n\n"
                "Table 1 summarizes the work breakdown schedule used for planning and tracking progress."
            ),
        ),
        (
            "Project Planning",
            "s70",
            (
                "Project planning followed an incremental delivery model: establish the API contract and "
                "database schema first, then parallelize mobile and admin clients against stable endpoints. "
                "Risk areas identified early included JWT session refresh on mobile, CORS configuration for "
                "LAN device testing, multipart upload limits for business event photos, and role separation "
                "between consumer, partner, and admin paths. Configuration is centralized in "
                "application.yml (datasource, JWT secret, upload directories, CORS patterns), mobile "
                "EXPO_PUBLIC_API_URL, and admin VITE_API_BASE_URL.\n\n"
                "[PLACEHOLDER: Figure 1.2 - Project Gantt Chart]\n"
                "Description: Gantt chart derived from Table 1 timeline showing analysis, backend, mobile, "
                "admin, and documentation phases with dependencies (API before client feature completion).\n\n"
                "Quality activities included Spring Boot unit and integration tests, TypeScript type-checking "
                "on mobile and admin, and manual end-to-end verification of booking, approval, and moderation "
                "flows described in docs/GRADUATION_PROJECT_HANDOFF.md and docs/MOBILE_AND_ADMIN_GUIDE.md."
            ),
        ),
        (
            "Planning of the Development Phases",
            "s70",
            (
                "Development phases were ordered to reduce integration risk:\n\n"
                "Phase 1 — Foundation: MySQL schema via JPA entities, auth register/login/refresh/logout, "
                "governorate and category seed data, and public catalog reads.\n\n"
                "Phase 2 — Consumer experience: mobile Explore, search, filters, event PDP, favorites, "
                "bookings API integration, ratings, and user reports.\n\n"
                "Phase 3 — Business partner: business profile application and submit, admin approval, "
                "ROLE_BUSINESS lifecycle, partner event CRUD, time slots, photo uploads, and booking status updates.\n\n"
                "Phase 4 — Administration: admin-web authentication, moderation queues, user management patches, "
                "and GET /admin/analytics/summary dashboard metrics.\n\n"
                "Phase 5 — Stabilization: removal of runtime mock catalog in favor of API-only data, PayPal "
                "Sandbox checkout wiring, production data cleanup, and graduation documentation."
            ),
        ),
        (
            "The Scope of the Work",
            "s70",
            (
                "The scope of work encompasses requirements engineering, design, implementation, and validation "
                "of the Vibook platform as a cohesive software product. Requirements were derived from the "
                "implemented codebase and verified against REST controllers, entity relationships, and mobile/admin "
                "route maps. Stakeholder and user needs—discovery, booking, partner self-service, and platform "
                "oversight—are addressed through concrete API and UI flows rather than hypothetical extensions.\n\n"
                "This section establishes why the documented requirements exist: Jordan-focused event discovery, "
                "governed business onboarding, and trust mechanisms (ratings and moderation). Detailed functional "
                "and non-functional requirements appear in subsequent chapters; stakeholder profiles and user "
                "environment descriptions are expanded in Chapter 3 of this template."
            ),
        ),
        (
            "Stakeholders",
            "s70",
            (
                "Key stakeholders for Vibook include:\n\n"
                "• End consumers and guests — browse and book events; require accurate catalog data, secure accounts, "
                "and reliable booking confirmation.\n\n"
                "• Business partners — submit BusinessProfile applications, manage events and bookings after approval; "
                "require transparent status and listing tools.\n\n"
                "• Platform administrators — operate admin-web; approve businesses, moderate events and reports, "
                "and monitor analytics.\n\n"
                "• Academic supervisor and evaluation committee — assess software engineering process, architecture, "
                "and documentation completeness.\n\n"
                "• Development team — implement and maintain backend, mobile, and admin-web codebases under version control.\n\n"
                "[PLACEHOLDER: Table 1.1 - Stakeholder Summary]\n"
                "Description: Tabular summary of each stakeholder type with representative role, responsibilities, "
                "success criteria, and involvement level in requirements and acceptance testing."
            ),
        ),
        (
            "Scope",
            "s70",
            (
                "In scope for this graduation project:\n"
                "• Spring Boot 3.3 REST API with Spring Security, JWT, JPA, and MySQL persistence.\n"
                "• Mobile app: authentication, Explore/search/filters, event PDP, bookings (including PayPal Sandbox), "
                "favorites, ratings, reports, business partner hub, bilingual UI (English/Arabic) with RTL support.\n"
                "• Admin-web: JWT login for ADMIN role, business approval, event and booking moderation, ratings and "
                "reports queues, analytics summary.\n"
                "• Static file serving for profile images and business logos/banners.\n\n"
                "Explicitly out of scope (not implemented as production backend features):\n"
                "• Premium membership billing, wallet balances, and voucher redemption beyond static mobile UI.\n"
                "• Legacy mobile PDP routes for restaurant, stay, experience, package, and organizer shells.\n"
                "• Push notification delivery infrastructure (notification list is presentation-only).\n\n"
                "Future work may extend payment providers beyond Sandbox, enrich notifications, and expand catalog "
                "types; such items are not claimed as delivered functionality in this report."
            ),
        ),
        (
            "Definitions, Acronyms, and Abbreviations",
            "Normal",
            (
                "API — Application Programming Interface; Vibook exposes REST endpoints under /api/v1.\n"
                "PDP — Product Detail Page; mobile route app/event/[id].tsx for a single event listing.\n"
                "JWT — JSON Web Token; access tokens (24h) and refresh tokens (30 days, stored server-side).\n"
                "JPA — Java Persistence API; Hibernate maps entities to MySQL tables.\n"
                "SPA — Single Page Application; admin-web built with Vite and React Router.\n"
                "Expo — Framework and toolchain for the React Native mobile application.\n"
                "DTO — Data Transfer Object; request/response records in backend/dto.\n"
                "CORS — Cross-Origin Resource Sharing; configured for localhost and LAN origins.\n"
                "RTL — Right-to-left layout for Arabic locale via I18nManager and RtlLayout.\n"
                "JOD — Jordanian Dinar; default currency display context in mobile locale settings.\n"
                "USER / ADMIN / ROLE_BUSINESS — Spring Security roles used in authorization rules.\n"
                "BusinessProfile — Partner application entity with approval workflow before event publishing.\n"
                "ModerationReport — User-submitted report queued for admin resolution."
            ),
        ),
        (
            "References",
            "Normal",
            (
                "[1] Spring Boot Documentation, version 3.3 — https://docs.spring.io/spring-boot/docs/3.3.x/reference/html/\n"
                "[2] Spring Security Reference — https://docs.spring.io/spring-security/reference/\n"
                "[3] Spring Data JPA Documentation — https://docs.spring.io/spring-data/jpa/docs/current/reference/html/\n"
                "[4] Hibernate ORM User Guide — https://docs.hibernate.org/orm/documentation/\n"
                "[5] MySQL 8.0 Reference Manual — https://dev.mysql.com/doc/refman/8.0/en/\n"
                "[6] JJWT (Java JWT) library, version 0.12.6 — https://github.com/jwtk/jjwt\n"
                "[7] Expo Documentation, SDK 54 — https://docs.expo.dev/\n"
                "[8] React Native Documentation, version 0.81 — https://reactnative.dev/docs/getting-started\n"
                "[9] expo-router Documentation — https://docs.expo.dev/router/introduction/\n"
                "[10] Zustand State Management — https://github.com/pmndrs/zustand\n"
                "[11] React 19 Documentation — https://react.dev/\n"
                "[12] Vite Documentation, version 8 — https://vite.dev/guide/\n"
                "[13] Axios HTTP Client — https://axios-http.com/docs/intro\n"
                "[14] React Router Documentation, version 7 — https://reactrouter.com/\n"
                "[15] Recharts Documentation — https://recharts.org/en-US/\n"
                "[16] PayPal REST API (Sandbox) — https://developer.paypal.com/api/rest/\n"
                "[17] Jakarta Bean Validation — https://jakarta.ee/specifications/bean-validation/\n"
                "[18] TypeScript Handbook — https://www.typescriptlang.org/docs/\n"
                "[19] RFC 7519 — JSON Web Token (JWT) — https://datatracker.ietf.org/doc/html/rfc7519\n"
                "[20] Vibook project repository — docs/GRADUATION_PROJECT_HANDOFF.md, docs/MOBILE_AND_ADMIN_GUIDE.md"
            ),
        ),
        (
            "Overview",
            "s70",
            (
                "This chapter introduces the Vibook graduation project: its purpose as a Jordan-focused event "
                "booking marketplace, the three-part architecture (mobile, backend API, admin console), and the "
                "planning structure used to deliver verified functionality. Subsequent chapters of this document "
                "detail positioning, stakeholder and user descriptions, product capabilities, constraints, and "
                "technical design artifacts including use cases, database schema, and diagrams.\n\n"
                "Readers should use Section 1.8 Scope together with the implementation handoff documentation "
                "to distinguish production API-backed features from presentation-only mobile screens. Figures "
                "and tables marked as placeholders in this chapter should be replaced with captures and diagrams "
                "from the running system and database before final submission."
            ),
        ),
        (
            "Product Scenarios",
            "s70",
            (
                "The following product scenario summarizes primary interactions between actors and the implemented "
                "Vibook system. Consumers and guests interact with the mobile app; business partners use mobile "
                "business routes after profile approval; administrators use admin-web; all privileged operations "
                "are enforced by the backend JWT authorization model."
            ),
        ),
    ]


def rebuild_chapter1_body(doc, intro_idx, chapter2_idx, style_refs):
    """Replace placeholder body between Introduction H1 and Chapter 2."""
    intro_p = doc.paragraphs[intro_idx]

    to_delete = []
    for i in range(intro_idx + 1, chapter2_idx):
        to_delete.append(doc.paragraphs[i])
    for p in reversed(to_delete):
        delete_paragraph(p)

    intro_p = doc.paragraphs[intro_idx]
    anchor = intro_p
    sections = build_sections()

    for _idx, (heading, body_style, body_text) in enumerate(sections):
        anchor = insert_paragraph_after(
            anchor, heading, "Heading 2", style_refs=style_refs
        )

        if heading == "Tasks":
            for para_part in body_text.split("\n\n"):
                anchor = insert_paragraph_after(
                    anchor, para_part, body_style, style_refs=style_refs
                )
            anchor = insert_paragraph_after(
                anchor, "Table 1 Project Timeline", "s75", style_refs=style_refs
            )
            anchor = insert_paragraph_after(
                anchor, "", "s3", style_refs=style_refs
            )
            cap = insert_paragraph_after(
                anchor, "Table 1 Project Timeline", "Caption", style_refs=style_refs
            )
            anchor = insert_paragraph_after(
                cap,
                "Table 1 lists the work breakdown structure, schedule, and completion status for major Vibook "
                "delivery tasks from requirements through documentation.",
                "s70",
                style_refs=style_refs,
            )
            continue

        for para_part in body_text.split("\n\n"):
            style = "InfoBlue" if para_part.strip().startswith("[PLACEHOLDER:") else body_style
            anchor = insert_paragraph_after(
                anchor, para_part, style, style_refs=style_refs
            )

    anchor = insert_paragraph_after(anchor, "", "Normal", style_refs=style_refs)
    anchor = insert_paragraph_after(
        anchor,
        "[PLACEHOLDER: Figure 1.3 - Vibook Use Case Diagram]\n"
        "Description: Show actors Guest, Consumer (USER), Business Partner, and Admin interacting with "
        "use cases Discover Events, Book Event, Manage Business Profile, Publish Event, Moderate Content, "
        "and View Analytics based on implemented mobile and admin flows.",
        "InfoBlue",
        style_refs=style_refs,
    )
    anchor = insert_paragraph_after(anchor, "", "Normal", style_refs=style_refs)
    anchor = insert_paragraph_after(
        anchor, "Figure 1.3 Vibook Use Case Diagram", "Caption", style_refs=style_refs
    )
    insert_paragraph_after(
        anchor,
        "Figure 1.3 illustrates primary interactions between Vibook actors and the implemented booking, "
        "partner onboarding, and administration capabilities. Replace the placeholder with the project-specific "
        "diagram before submission.",
        "InfoBlue",
        style_refs=style_refs,
    )


def validate_docx(path):
    """Return list of validation issues for a saved .docx."""
    issues = []
    path = Path(path)
    if not path.is_file():
        return [f"Missing file: {path}"]

    import zipfile
    import xml.etree.ElementTree as ET

    try:
        with zipfile.ZipFile(path) as zf:
            bad = zf.testzip()
            if bad:
                issues.append(f"Zip integrity failure at {bad}")
            for name in zf.namelist():
                if name.endswith((".xml", ".rels")):
                    try:
                        ET.fromstring(zf.read(name))
                    except ET.ParseError as exc:
                        issues.append(f"XML parse error in {name}: {exc}")
            xml = zf.read("word/document.xml").decode("utf-8")
    except zipfile.BadZipFile:
        return [f"Not a valid zip/docx archive: {path}"]

    no_tbl = re.sub(r"<w:tbl>.*?</w:tbl>", "", xml, flags=re.S)
    bare_p = len(re.findall(r"<w:p>(?!\s*<w:pPr)", no_tbl))
    if bare_p:
        issues.append(
            f"word/document.xml body contains {bare_p} paragraph(s) without w:pPr (Word may refuse to open)"
        )
    if xml.count("bookmarkStart") != xml.count("bookmarkEnd"):
        issues.append("Unbalanced bookmarkStart/bookmarkEnd in document.xml")
    return issues


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--source",
        type=Path,
        default=DEFAULT_SOURCE,
        help="Pristine graduation template .docx",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help="Output .docx path (default: Vibook_Introduction_FIXED.docx)",
    )
    args = parser.parse_args()

    if not args.source.is_file():
        raise SystemExit(f"Source template not found: {args.source}")

    work_copy = args.output.with_suffix(".work.docx")
    shutil.copy2(args.source, work_copy)

    doc = Document(str(work_copy))
    style_refs = _style_ref_cache(doc)
    ch1, ch2 = find_chapter1_bounds(doc)
    if ch1 is None or ch2 is None:
        raise SystemExit("Could not locate Chapter 1 / Chapter 2 boundaries")

    intro_idx = None
    for i in range(ch1, ch2):
        if doc.paragraphs[i].text.strip() == "Introduction":
            intro_idx = i
            break
    if intro_idx is None:
        raise SystemExit("Introduction heading not found")

    update_timeline_table(doc)
    rebuild_chapter1_body(doc, intro_idx, ch2, style_refs)
    move_timeline_table_after_tasks_label(doc)
    doc.save(str(args.output))
    work_copy.unlink(missing_ok=True)

    issues = validate_docx(args.output)
    if issues:
        print("Validation issues:", file=sys.stderr)
        for issue in issues:
            print(f"  - {issue}", file=sys.stderr)
        raise SystemExit(1)

    print(f"Saved {args.output} (validated OK)")


def move_timeline_table_after_tasks_label(doc):
    """Word may leave table 0 at end of chapter after paragraph rebuild."""
    if not doc.tables:
        return
    table = doc.tables[0]
    tbl = table._tbl
    parent = tbl.getparent()
    parent.remove(tbl)
    anchor = None
    seen_tasks = False
    for p in doc.paragraphs:
        if p.text.strip() == "Tasks" and p.style.name == "Heading 2":
            seen_tasks = True
        if seen_tasks and p.text.strip() == "Table 1 Project Timeline" and p.style.name == "s75":
            anchor = p
            break
    if anchor is not None:
        anchor._p.addnext(tbl)
    print("Introduction sections 1.1–1.11 filled.")


if __name__ == "__main__":
    main()
